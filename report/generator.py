"""
report/generator.py — Word report generator
"""

import os, json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colours
DARK_BLUE  = "1F3864"
MID_BLUE   = "2E75B6"
LIGHT_BLUE = "D5E8F0"
GREEN      = "E2EFDA"
ORANGE     = "FCE4D6"
RED        = "FFE0E0"
YELLOW     = "FFF2CC"
GREY       = "F2F2F2"
WHITE      = "FFFFFF"

STATUS_COLOR = {
    "EXCELLENT":    GREEN,
    "GOOD":         GREEN,
    "ACCEPTABLE":   YELLOW,
    "PASS":         GREEN,
    "POOR":         ORANGE,
    "FAIL":         RED,
    "CRITICAL":     RED,
    "HIGH":         ORANGE,
    "MEDIUM":       YELLOW,
    "LOW":          GREEN,
    "MONITOR":      YELLOW,
    "UNKNOWN":      GREY,
    "STRESS ONLY":  YELLOW,
}

# Text colours matching priority badges (used in _render_rec_table badge cell)
PRIORITY_TEXT_COLOR = {
    "CRITICAL": "C00000",
    "HIGH":     "E26B0A",
    "MEDIUM":   "BF8F00",
    "LOW":      "375623",
}

# ─── helpers ──────────────────────────────────────────────────────────────────
def _shd(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), hex_color)
    tcPr.append(s)

def _ct(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)

def _h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(WHITE)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), DARK_BLUE)
    pPr.append(s)

def _h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(MID_BLUE)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), MID_BLUE)
    pBdr.append(bot); pPr.append(pBdr)

def _body(doc, text, size=10):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)

def _obs(doc, icon, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{icon}  {text}"); r.font.size = Pt(10)

def _table(doc, headers, rows, header_bg=MID_BLUE):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        _ct(tbl.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(tbl.rows[0].cells[i], header_bg)
    for ri, row in enumerate(rows):
        tr = tbl.add_row(); bg = GREY if ri % 2 else WHITE
        for ci, v in enumerate(row):
            _ct(tr.cells[ci], str(v)); _shd(tr.cells[ci], bg)
    return tbl

def _status_row(tbl, values, status_col_idx, row_bg=WHITE):
    tr = tbl.add_row()
    for ci, v in enumerate(values):
        bold = ci == status_col_idx
        _ct(tr.cells[ci], str(v), bold=bold,
            align=WD_ALIGN_PARAGRAPH.CENTER if ci == status_col_idx else WD_ALIGN_PARAGRAPH.LEFT)
        color = STATUS_COLOR.get(str(v).upper().split()[0], row_bg) if ci == status_col_idx else row_bg
        _shd(tr.cells[ci], color)


# ─── Main builder ─────────────────────────────────────────────────────────────
def build_report(all_data: dict, llm_analysis: str, output_dir: str,
                 llm_structured: dict = None) -> str:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    device   = all_data.get("device", {})
    suit     = all_data.get("suitability", {})
    cold     = all_data.get("cold_start", {})
    warm     = all_data.get("warm_start", {})
    snaps    = all_data.get("snapshots", [])
    battery  = all_data.get("battery", {})
    network  = all_data.get("network", {})
    app_info = all_data.get("app_info", {})
    now      = datetime.now().strftime("%d %B %Y %H:%M")
    ram_gb   = device.get("ram_total_gb", 0)

    # ── Title Page ───────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Android Mobile Performance Report")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(app_info.get("app_name", app_info.get("package", "Unknown App")))
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(MID_BLUE)

    doc.add_paragraph()
    tbl = doc.add_table(rows=7, cols=2); tbl.style = "Table Grid"
    meta = [
        ("Application Package",  app_info.get("package", "N/A")),
        ("App Version",          app_info.get("version", "N/A")),
        ("Test Date",            now),
        ("Device",               f"{device.get('brand','')} {device.get('model','')}"),
        ("Android",              f"{device.get('android_ver','')} (API {device.get('api_level','')})"),
        ("Session Duration",     all_data.get("session_duration_min", "N/A")),
        ("Prepared By",          "Performance Engineering CoE"),
    ]
    for i, (k, v) in enumerate(meta):
        _shd(tbl.rows[i].cells[0], LIGHT_BLUE); _ct(tbl.rows[i].cells[0], k, bold=True)
        _ct(tbl.rows[i].cells[1], v)
    doc.add_page_break()

    # ── Section 1 — Device ───────────────────────────────────────────────
    _h1(doc, "1.  Device Specifications & Suitability Assessment")
    _h2(doc, "1.1  Hardware Profile")
    _table(doc, ["Parameter", "Value"], [
        ("Manufacturer",     f"{device.get('manufacturer','')}"),
        ("Model",            f"{device.get('brand','')} {device.get('model','')}"),
        ("Android Version",  f"{device.get('android_ver','')} (API {device.get('api_level','')})"),
        ("CPU",              f"{device.get('cpu_cores','')} cores @ {device.get('cpu_max_freq_ghz','')} GHz ({device.get('hardware','')})"),
        ("RAM",              f"{ram_gb} GB total"),
        ("CPU Architecture", device.get("cpu_abi", "N/A")),
        ("Screen",           f"{device.get('screen_size','')} @ {device.get('screen_density','')} dpi"),
        ("Serial",           device.get("serial", "N/A")),
    ])

    doc.add_paragraph()
    _h2(doc, "1.2  Device Suitability Assessment")
    score = suit.get("score", 0)
    tier  = suit.get("tier", "UNKNOWN")
    _body(doc, f"Suitability Score: {score}/100 — {tier}")
    details = suit.get("details", {})
    suit_rows = [
        ("RAM",              f"{ram_gb} GB",   ">= 4 GB",  details.get("ram", {}).get("status", "?")),
        ("Android API",      device.get("api_level",""),">= 26 (Android 8)",details.get("api",{}).get("status","?")),
        ("CPU Cores",        device.get("cpu_cores",""),">= 4 cores", details.get("cores",{}).get("status","?")),
        ("CPU Frequency",    f"{device.get('cpu_max_freq_ghz','')} GHz",">= 1.8 GHz",details.get("cpu_freq",{}).get("status","?")),
    ]
    tbl2 = doc.add_table(rows=1, cols=4); tbl2.style = "Table Grid"
    for i, h in enumerate(["Component", "Device Value", "Min Recommended", "Status"]):
        _ct(tbl2.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(tbl2.rows[0].cells[i], MID_BLUE)
    for ri, row in enumerate(suit_rows):
        _status_row(tbl2, row, 3, GREY if ri%2 else WHITE)

    if suit.get("issues"):
        doc.add_paragraph()
        for iss in suit["issues"]:
            _obs(doc, "⚠️", iss)

    # App RAM usage context
    if ram_gb > 0 and snaps:
        first_pss = snaps[0].get("parsed", {}).get("total_pss_mb", 0)
        pct = round((first_pss / (ram_gb * 1024)) * 100, 2) if first_pss else 0
        doc.add_paragraph()
        _obs(doc, "📊", f"App uses {first_pss} MB PSS at idle = {pct}% of device's {ram_gb} GB RAM. "
                        f"{'Well within safe range.' if pct < 5 else 'Monitor for low-memory kills on devices with less RAM.'}")

    doc.add_page_break()

    # ── Section 2 — Start Times ──────────────────────────────────────────
    _h1(doc, "2.  App Start Time — Cold Start & Warm Start")
    _h2(doc, "2.1  Results")
    tbl3 = doc.add_table(rows=1, cols=5); tbl3.style = "Table Grid"
    for i, h in enumerate(["Type", "Run 1 (ms)", "Run 2 (ms)", "Run 3 (ms)", "Average (ms)"]):
        _ct(tbl3.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(tbl3.rows[0].cells[i], MID_BLUE)
    for st, data in [("Cold Start", cold), ("Warm Start", warm)]:
        vals = data.get("values", [])
        tr = tbl3.add_row()
        _ct(tr.cells[0], st, bold=True); _shd(tr.cells[0], GREY)
        for ci in range(1, 4):
            v = str(vals[ci-1]) if ci-1 < len(vals) else "N/A"
            _ct(tr.cells[ci], v, align=WD_ALIGN_PARAGRAPH.CENTER); _shd(tr.cells[ci], WHITE)
        _ct(tr.cells[4], str(data.get("avg", "N/A")), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(tr.cells[4], LIGHT_BLUE)

    doc.add_paragraph()
    _h2(doc, "2.2  Benchmark Comparison")
    tbl4 = doc.add_table(rows=1, cols=4); tbl4.style = "Table Grid"
    for i, h in enumerate(["Metric", "Measured", "Google Standard", "Result"]):
        _ct(tbl4.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(tbl4.rows[0].cells[i], MID_BLUE)
    cold_avg = cold.get("avg", 0); warm_avg = warm.get("avg", 0)
    cold_status = "EXCELLENT" if cold_avg<=1000 else "GOOD" if cold_avg<=2000 else "ACCEPTABLE" if cold_avg<=5000 else "POOR"
    warm_status = "EXCELLENT" if warm_avg<=200  else "GOOD" if warm_avg<=800  else "ACCEPTABLE" if warm_avg<=2000 else "POOR"
    _status_row(tbl4, ["Cold Start", f"{cold_avg} ms", "< 2000 ms (Good)", cold_status], 3)
    _status_row(tbl4, ["Warm Start", f"{warm_avg} ms", "< 800 ms (Good)",  warm_status], 3, GREY)
    doc.add_page_break()

    # ── Section 3 — Memory ───────────────────────────────────────────────
    _h1(doc, "3.  Memory Analysis")

    if snaps:
        _h2(doc, "3.1  Memory Snapshot Trend")
        tbl5 = doc.add_table(rows=1, cols=6); tbl5.style = "Table Grid"
        for i, h in enumerate(["Snapshot", "Time", "Total PSS (MB)", "Native Heap (MB)", "GL mtrack (MB)", "CPU %"]):
            _ct(tbl5.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shd(tbl5.rows[0].cells[i], MID_BLUE)
        for ri, s in enumerate(snaps):
            p = s.get("parsed", {})
            tr = tbl5.add_row(); bg = GREY if ri%2 else WHITE
            vals = [s["label"], s["ts"],
                    p.get("total_pss_mb","?"), p.get("native_heap_mb","?"),
                    p.get("gl_mtrack_mb","?"), s.get("cpu_pct","?")]
            for ci, v in enumerate(vals):
                _ct(tr.cells[ci], str(v)); _shd(tr.cells[ci], bg)

        # Growth analysis (first vs last)
        if len(snaps) >= 2:
            doc.add_paragraph()
            _h2(doc, "3.2  Memory Growth Analysis")
            first_p = snaps[0].get("parsed", {}); last_p = snaps[-1].get("parsed", {})
            def _growth(key):
                a = first_p.get(key, 0) or 0; b = last_p.get(key, 0) or 0
                pct = round(((b-a)/a)*100, 1) if a else 0
                return a, b, pct
            tbl6 = doc.add_table(rows=1, cols=5); tbl6.style = "Table Grid"
            for i, h in enumerate(["Component", "First Snapshot", "Last Snapshot", "Growth %", "Concern"]):
                _ct(tbl6.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
                _shd(tbl6.rows[0].cells[i], MID_BLUE)
            growth_items = [
                ("Total PSS",    "total_pss_mb",    60,  "Stable growth expected"),
                ("Native Heap",  "native_heap_mb",  100, "Watch for leaks if >50% growth"),
                ("GL mtrack",    "gl_mtrack_mb",    150, "Camera textures — release after capture"),
            ]
            for ri, (label, key, warn_pct, note_text) in enumerate(growth_items):
                a, b, pct = _growth(key)
                concern_color = RED if pct > warn_pct else (YELLOW if pct > warn_pct/2 else GREEN)
                concern = "CRITICAL" if pct > warn_pct else ("MONITOR" if pct > warn_pct/2 else "GOOD")
                tr = tbl6.add_row(); bg = GREY if ri%2 else WHITE
                _ct(tr.cells[0], label, bold=True); _shd(tr.cells[0], bg)
                _ct(tr.cells[1], f"{a} MB");        _shd(tr.cells[1], bg)
                _ct(tr.cells[2], f"{b} MB");        _shd(tr.cells[2], bg)
                _ct(tr.cells[3], f"+{pct}%");       _shd(tr.cells[3], bg)
                _ct(tr.cells[4], concern, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                _shd(tr.cells[4], concern_color)

        doc.add_paragraph()
        _h2(doc, "3.3  Benchmark Comparison")
        idle_pss = snaps[0].get("parsed", {}).get("total_pss_mb", 0) if snaps else 0
        gl_idle  = snaps[0].get("parsed", {}).get("gl_mtrack_mb", 0) if snaps else 0
        swap     = snaps[0].get("parsed", {}).get("swap_pss_mb", 0) if snaps else 0
        idle_status  = "EXCELLENT" if idle_pss<=50 else "GOOD" if idle_pss<=100 else "ACCEPTABLE" if idle_pss<=150 else "POOR"
        gl_status    = "EXCELLENT" if gl_idle<=10  else "GOOD" if gl_idle<=20  else "ACCEPTABLE" if gl_idle<=40  else "POOR"
        swap_status  = "EXCELLENT" if swap==0 else "GOOD" if swap<=10 else "ACCEPTABLE" if swap<=50 else "POOR"
        tbl7 = doc.add_table(rows=1, cols=4); tbl7.style = "Table Grid"
        for i, h in enumerate(["Metric", "Measured", "Google Standard", "Status"]):
            _ct(tbl7.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shd(tbl7.rows[0].cells[i], MID_BLUE)
        _status_row(tbl7, ["Idle PSS",    f"{idle_pss} MB", "< 100 MB",      idle_status], 3)
        _status_row(tbl7, ["GL mtrack",   f"{gl_idle} MB",  "< 20 MB",       gl_status],   3, GREY)
        _status_row(tbl7, ["Swap PSS",    f"{swap} MB",     "0 MB (optimal)",swap_status],  3)

    doc.add_page_break()

    # ── Section 4 — CPU ──────────────────────────────────────────────────
    _h1(doc, "4.  CPU Usage Analysis")
    if snaps:
        cpu_vals = [s.get("cpu_pct", 0) for s in snaps if s.get("cpu_pct", 0) > 0]
        idle_cpu = snaps[0].get("cpu_pct", 0) if snaps else 0
        peak_cpu = max(cpu_vals) if cpu_vals else 0
        avg_cpu  = round(sum(cpu_vals)/len(cpu_vals), 1) if cpu_vals else 0

        _h2(doc, "4.1  CPU Summary")
        tbl8 = doc.add_table(rows=1, cols=4); tbl8.style = "Table Grid"
        for i, h in enumerate(["Metric", "Measured", "Google Standard", "Status"]):
            _ct(tbl8.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shd(tbl8.rows[0].cells[i], MID_BLUE)
        idle_s = "EXCELLENT" if idle_cpu<=2 else "GOOD" if idle_cpu<=5 else "ACCEPTABLE" if idle_cpu<=10 else "POOR"
        avg_s  = "EXCELLENT" if avg_cpu<=30  else "GOOD" if avg_cpu<=50  else "ACCEPTABLE" if avg_cpu<=80  else "POOR"
        _status_row(tbl8, ["Idle CPU",    f"{idle_cpu}%", "< 5%",   idle_s], 3)
        _status_row(tbl8, ["Avg Active",  f"{avg_cpu}%",  "< 50%",  avg_s],  3, GREY)
        _status_row(tbl8, ["Peak CPU",    f"{peak_cpu}%", "< 200%", "ACCEPTABLE" if peak_cpu<=200 else "MONITOR"], 3)

        # % of device capacity
        cores = device.get("cpu_cores", 1) or 1
        max_pct = cores * 100
        if peak_cpu > 0:
            doc.add_paragraph()
            _obs(doc, "📊", f"Peak CPU {peak_cpu}% of single-core 100% scale. "
                            f"Device has {cores} cores (max {max_pct}%). "
                            f"Peak used {round(peak_cpu/max_pct*100,1)}% of total device CPU capacity.")

    doc.add_page_break()

    # ── Section 5 — Battery / GPS / Camera ──────────────────────────────
    _h1(doc, "5.  Battery, GPS & Camera Hardware Analysis")
    was_charging = battery.get("was_charging", False)
    if was_charging:
        _obs(doc, "⚠️", "Device was CHARGING via USB during this test session. "
                        "Battery drain figures are not available. Re-run with USB unplugged for drain metrics.")

    _h2(doc, "5.1  Session Overview")
    _table(doc, ["Metric", "Value", "Notes"], [
        ("Session Duration",    battery.get("session_duration","N/A"),      "From batterystats reset to capture"),
        ("Charging During Test",str(was_charging),                           "USB connected = cannot measure drain"),
        ("Battery Capacity",    f"{battery.get('battery_capacity_mah',0)} mAh", "Device total capacity"),
        ("Max Temperature",     f"{battery.get('temp_max_c','N/A')} deg C",  "Normal < 45 deg C"),
        ("WorkManager Jobs",    battery.get("workmanager_jobs",0),           "Background job executions"),
        ("Firebase Jobs",       battery.get("firebase_jobs",0),              "Data transport flush events"),
        ("WiFi Radio Wakeups",  battery.get("wifi_wakeups",0),               "App-triggered WiFi wakeups"),
    ])

    doc.add_paragraph()
    _h2(doc, "5.2  Camera Hardware Usage")
    cam_sessions = battery.get("camera_sessions", 0)
    cam_sec      = battery.get("camera_total_sec", 0)
    _table(doc, ["Metric", "Value", "Benchmark / Note"], [
        ("Total Camera Sessions",    cam_sessions,          "Hardware open/close events"),
        ("Total Camera Active Time", f"{cam_sec} sec ({round(cam_sec/60,1)} min)", "Camera sensor powered on"),
        ("Avg Session Duration",     f"{round(cam_sec/cam_sessions,1) if cam_sessions else 0} sec", "< 5 sec = preview/init cycles"),
        ("Double-Open Pattern",      "Check if sessions appear in pairs", "App may open camera twice per photo"),
    ])

    doc.add_paragraph()
    _h2(doc, "5.3  GPS Hardware Usage")
    gps_act = battery.get("gps_activations", 0)
    gps_sec = battery.get("gps_total_sec", 0)
    dur_min = _parse_duration_to_min(battery.get("session_duration",""))
    gps_per_min = round(gps_act / dur_min, 1) if dur_min > 0 else 0
    gps_status = "EXCELLENT" if gps_per_min<=2 else "GOOD" if gps_per_min<=6 else "ACCEPTABLE" if gps_per_min<=12 else "POOR"
    _table(doc, ["Metric", "Value", "Google Standard"], [
        ("Total GPS Activations",    gps_act,                f"< 6/min recommended (yours: {gps_per_min}/min)"),
        ("GPS Active Time",          f"{gps_sec} sec",       "Total time GPS radio was powered on"),
        ("GPS Frequency",            f"{gps_per_min}/min",   "Navigation: 1/sec, Field work: 1/10sec"),
        ("GPS Status",               gps_status,             "Based on activation frequency"),
    ])
    doc.add_page_break()

    # ── Section 6 — Network ──────────────────────────────────────────────
    _h1(doc, "6.  Network & SDK Analysis")
    sdks = network.get("sdks_detected", [])
    _h2(doc, "6.1  SDKs & Services Detected")
    if sdks:
        _table(doc, ["SDK / Service", "Detected"], [(s, "Yes") for s in sdks])
    else:
        _body(doc, "No third-party SDKs detected in logs.")

    redundant = network.get("redundant_calls", [])
    if redundant:
        doc.add_paragraph()
        _h2(doc, "6.2  Redundant / Duplicate Calls")
        for rc in redundant:
            _table(doc, ["Field", "Detail"], [
                ("Type",       rc.get("type","")),
                ("Occurrences",rc.get("occurrences","")),
                ("Evidence",   rc.get("evidence","")),
                ("Severity",   rc.get("severity","")),
                ("Fix",        rc.get("fix","")),
            ])
            doc.add_paragraph()

    doc.add_paragraph()
    _h2(doc, "6.3  Background Job Activity")
    top_jobs = network.get("top_jobs", [])
    if top_jobs:
        _table(doc, ["Job Service", "Executions"], [(j[0].split("/")[-1], j[1]) for j in top_jobs])

    doc.add_page_break()

    # ── Section 7 — LLM Analysis + Colour-Coded Recommendations ────────
    _h1(doc, "7.  AI-Powered Performance Analysis & Recommendations")
    _body(doc, "Analysis generated by LLM based on all collected data. Benchmark comparisons use Google Android Vitals standards.")
    doc.add_paragraph()

    if llm_analysis:
        # ── Parse LLM text into structured sections ───────────────────
        sections = _parse_llm_sections(llm_analysis)

        for sec_title, sec_lines in sections.items():
            is_rec = any(k in sec_title.upper() for k in
                         ["RECOMMEND", "PRIORITY", "FINDING", "ACTION"])

            # Section header
            p = doc.add_paragraph()
            r = p.add_run(sec_title)
            r.bold = True; r.font.size = Pt(12)
            r.font.color.rgb = RGBColor.from_string(MID_BLUE)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)

            if is_rec:
                # Build colour-coded recommendation table
                _build_rec_table(doc, sec_lines)
            else:
                for line in sec_lines:
                    if line.strip():
                        p2 = doc.add_paragraph()
                        p2.paragraph_format.left_indent = Cm(0.3)
                        p2.paragraph_format.space_after = Pt(3)
                        r2 = p2.add_run(line.strip())
                        r2.font.size = Pt(10)
    else:
        _body(doc, "[LLM analysis not available — check API key configuration]")

    doc.add_page_break()

    # ── Section 8 — Prioritised Recommendations (Standalone) ────────────
    _h1(doc, "8.  Prioritised Recommendations")

    # Determine source label for the intro line
    llm_structured = llm_structured or {}
    llm_recs_raw   = llm_structured.get("recommendations", [])
    rule_recs      = _build_recommendations(battery, network, snaps, cold, warm)
    merged_recs    = _merge_recommendations(llm_recs_raw, rule_recs)

    if llm_recs_raw:
        source_note = (
            "Recommendations generated by AI analysis of all collected metrics, "
            "merged with rule-based benchmark checks. "
            "Colour coding reflects urgency based on actual usage pattern (1-2 trips/day)."
        )
    else:
        source_note = (
            "All findings from this test session ranked by real-world impact. "
            "Colour coding reflects urgency based on actual usage pattern (1-2 trips/day). "
            "Note: LLM analysis was not available — recommendations are rule-based only."
        )
    _body(doc, source_note)

    # Overall risk badge from LLM if available
    overall_risk = llm_structured.get("overall_risk", "")
    if overall_risk:
        cfg = PRIORITY_CONFIG.get(overall_risk, PRIORITY_CONFIG["MEDIUM"])
        p = doc.add_paragraph()
        r = p.add_run(f"  Overall Risk Level: {cfg[2]}  ")
        r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string(cfg[1])
        pPr = p._p.get_or_add_pPr()
        s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear")
        s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), cfg[0])
        pPr.append(s)

    doc.add_paragraph()

    # Legend
    _h2(doc, "Colour Legend")
    leg_tbl = doc.add_table(rows=1, cols=4); leg_tbl.style = "Table Grid"
    for i, (label, color, desc) in enumerate([
        ("🔴  CRITICAL — MUST FIX",  "C00000", "Immediate action. Affects all users every shift."),
        ("🟠  HIGH — Should Fix",     "E26B0A", "Fix in next sprint. Noticeable user impact."),
        ("🟡  MEDIUM — Nice to Fix",  "BF8F00", "Fix when capacity allows. Minor impact."),
        ("🟢  LOW — Backlog",         "375623", "Cleanup item. Minimal real-world impact."),
    ]):
        c = leg_tbl.rows[0].cells[i]
        ct_legend(c, label, color)
    doc.add_paragraph()

    _render_rec_table(doc, merged_recs)

    doc.add_page_break()

    # ── Section 9 — Battery Testing Guidance ────────────────────────────
    _h1(doc, "9.  Battery Testing — Approach & Methodology")

    was_charging = battery.get("was_charging", False)
    if was_charging:
        p = doc.add_paragraph()
        r = p.add_run("⚠️  Device was charging via USB during this session — accurate battery drain could not be measured.")
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string("C00000")

    _body(doc, "Battery drain cannot be measured while the device is connected via USB because Android switches "
               "to a charging-optimised power state. The following approaches are recommended.")

    doc.add_paragraph()
    _h2(doc, "9.1  Recommended Approach — Wireless ADB (Best)")
    _table(doc, ["Step", "Action", "Command / Detail"], [
        ("1", "Enable Wireless Debugging on device",
              "Settings > Developer Options > Wireless Debugging > Enable"),
        ("2", "Note the device IP address",
              "Settings > About Phone > IP Address  OR  adb shell ip route"),
        ("3", "Connect wirelessly",
              "adb connect <device-ip>:5555"),
        ("4", "Unplug USB cable",
              "Device now runs on battery. adb stays connected over WiFi."),
        ("5", "Reset battery stats",
              "adb shell dumpsys batterystats --reset"),
        ("6", "Run your test session",
              "All framework commands work identically over wireless adb"),
        ("7", "Capture battery report",
              "adb shell dumpsys batterystats > battery_stats.txt"),
        ("8", "Capture bugreport",
              "adb bugreport bugreport.zip"),
    ])

    doc.add_paragraph()
    _h2(doc, "9.2  Alternative — Manual Reconnect Approach")
    _table(doc, ["Step", "Action", "Notes"], [
        ("1", "Reset stats (USB connected)",   "adb shell dumpsys batterystats --reset"),
        ("2", "Note battery % on device",      "e.g. start at 85% — check Settings > Battery"),
        ("3", "Unplug USB cable",               "Important: do this AFTER reset"),
        ("4", "Run test trips manually",         "Complete 1-2 full trips on device"),
        ("5", "Note battery % after trips",     "e.g. ended at 81% = 4% drain"),
        ("6", "Reconnect USB immediately",       "Reconnect before battery stats expire"),
        ("7", "Capture data within 30 seconds",  "adb bugreport bugreport.zip"),
    ])
    _obs(doc, "⚠️", "You have a short window after reconnecting USB — capture bugreport first before anything else.")

    doc.add_paragraph()
    _h2(doc, "9.3  What Accurate Battery Data Tells You")
    _table(doc, ["Metric", "How to Get It", "Why It Matters"], [
        ("mAh drain per trip",
         "charge_start - charge_end in batterystats",
         "Tells you exactly how much battery 1 trip costs"),
        ("Screen-on discharge rate",
         "'Screen on discharge: X mAh' in batterystats",
         "How fast battery drains with screen on during active use"),
        ("GPS mAh contribution",
         "Battery Historian — GPS rail usage",
         "Confirms if GPS fix reduces drain after optimisation"),
        ("Camera mAh contribution",
         "Battery Historian — Camera rail usage",
         "Quantifies camera power cost per photo"),
        ("Wakelock drain",
         "Battery Historian — wakelocks section",
         "Identifies background processes keeping CPU awake"),
        ("App vs system ratio",
         "dumpsys batterystats — per-app section",
         "What % of total drain is your app vs system services"),
        ("Projected full-shift drain",
         "Extrapolate from 1-hour sample",
         "Answers: will the device last a 12-hour shift?"),
    ])

    doc.add_paragraph()
    _h2(doc, "9.4  Battery Historian — Viewing the Report")
    _table(doc, ["Option", "How to Use", "Availability"], [
        ("bathist.ef.lc (Online)",
         "Upload bugreport.zip at https://bathist.ef.lc",
         "Free, no install — but can go offline"),
        ("Docker (Local)",
         "docker run -p 9999:9999 gcr.io/android-battery-historian/stable:3.0 --port 9999\nOpen http://localhost:9999",
         "Requires Docker Desktop — works offline"),
        ("Manual parsing (this framework)",
         "battery.py parses batterystats.txt directly for GPS, Camera, jobs",
         "Always available — built into this framework"),
        ("adb shell (Quick check)",
         "adb shell dumpsys batterystats | grep -E 'discharge|screen on|GPS'",
         "Quick summary without Historian"),
    ])
    _obs(doc, "💡", "The framework's battery.py module already extracts GPS activations, Camera sessions, "
                    "temperature, WorkManager and Firebase job counts directly from batterystats.txt — "
                    "no Battery Historian needed for these metrics.")
    _obs(doc, "📋", "ACTION: Re-run this test with USB unplugged (use wireless adb or manual reconnect method above) "
                    "to capture accurate mAh drain figures for the next report.")

    doc.add_page_break()

    # ── Section 10 — Scorecard ───────────────────────────────────────────
    _h1(doc, "10.  Performance Scorecard")
    cold_avg  = cold.get("avg", 0)
    warm_avg  = warm.get("avg", 0)
    cold_s2 = "EXCELLENT" if cold_avg<=1000 else "GOOD" if cold_avg<=2000 else "ACCEPTABLE" if cold_avg<=5000 else "POOR"
    warm_s2 = "EXCELLENT" if warm_avg<=200  else "GOOD" if warm_avg<=800  else "ACCEPTABLE" if warm_avg<=2000 else "POOR"
    idle_pss2 = snaps[0].get("parsed",{}).get("total_pss_mb",0) if snaps else 0

    scorecard = [
        ("Cold Start",    f"{cold.get('avg',0)} ms", "< 2000 ms",   cold_s2),
        ("Warm Start",    f"{warm.get('avg',0)} ms", "< 800 ms",    warm_s2),
        ("Idle Memory",   f"{idle_pss2} MB",         "< 100 MB",    "EXCELLENT" if idle_pss2<=50 else "GOOD" if idle_pss2<=100 else "POOR"),
        ("Idle CPU",      f"{snaps[0].get('cpu_pct',0) if snaps else '?'}%", "< 5%", "GOOD"),
        ("GPS Frequency", f"{gps_per_min}/min",      "< 6/min",     gps_status),
        ("Device Suitability", f"{suit.get('score',0)}/100", "> 70", suit.get("tier","?")),
        ("Swap Memory",   f"{snaps[0].get('parsed',{}).get('swap_pss_mb',0) if snaps else 0} MB", "0 MB", "EXCELLENT"),
    ]
    tbl_sc = doc.add_table(rows=1, cols=4); tbl_sc.style = "Table Grid"
    for i, h in enumerate(["Area", "Measured", "Standard", "Result"]):
        _ct(tbl_sc.rows[0].cells[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(tbl_sc.rows[0].cells[i], MID_BLUE)
    for ri, row in enumerate(scorecard):
        _status_row(tbl_sc, row, 3, GREY if ri%2 else WHITE)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"Performance Engineering CoE  |  Generated: {now}  |  Package: {app_info.get('package','N/A')}")
    r.font.size = Pt(8); r.italic = True; r.font.color.rgb = RGBColor.from_string("888888")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    safe_name = app_info.get("package","app").split(".")[-1]
    filename  = f"PerfReport_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    out_path  = os.path.join(output_dir, filename)
    doc.save(out_path)
    return out_path


def _parse_duration_to_min(duration_str):
    import re
    if not duration_str: return 0
    h = re.search(r"(\d+)h", duration_str); m = re.search(r"(\d+)m", duration_str)
    return (int(h.group(1))*60 if h else 0) + (int(m.group(1)) if m else 0)


# ─── Recommendation helpers ───────────────────────────────────────────────────

# Priority config: label, background hex, text hex, badge text
PRIORITY_CONFIG = {
    "CRITICAL": ("FFE0E0", "C00000", "🔴  CRITICAL — MUST FIX"),
    "HIGH":     ("FCE4D6", "E26B0A", "🟠  HIGH — Should Fix"),
    "MEDIUM":   ("FFF2CC", "BF8F00", "🟡  MEDIUM — Nice to Fix"),
    "LOW":      ("E2EFDA", "375623", "🟢  LOW — Backlog"),
}

def ct_legend(cell, text, hex_color):
    """Render a legend cell with the given background and white/dark text."""
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9)
    # Use dark text on light backgrounds (MEDIUM/LOW), white on dark (CRITICAL/HIGH)
    text_color = hex_color  # same as bg — we'll override
    r.font.color.rgb = RGBColor.from_string(hex_color)
    _shd(cell, {"C00000": "FFE0E0", "E26B0A": "FCE4D6",
                "BF8F00": "FFF2CC", "375623": "E2EFDA"}.get(hex_color, "F2F2F2"))


def _parse_llm_sections(text):
    """Split LLM text into sections by numbered headings."""
    import re
    sections = {}
    current_title = "Analysis"
    current_lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        # Detect section header: "1. Title" or "## Title" or "**Title**"
        is_header = (
            re.match(r"^\d+\.\s+[A-Z]", stripped) or
            stripped.startswith("##") or
            (stripped.startswith("**") and stripped.endswith("**") and len(stripped) < 60)
        )
        if is_header and current_lines:
            sections[current_title] = current_lines
            current_lines = []
            current_title = stripped.lstrip("#").strip().strip("*").strip()
        elif is_header:
            current_title = stripped.lstrip("#").strip().strip("*").strip()
        else:
            current_lines.append(stripped)
    if current_lines:
        sections[current_title] = current_lines
    return sections


def _detect_priority(line):
    """Detect priority level from a recommendation line."""
    upper = line.upper()
    if any(k in upper for k in ["CRITICAL", "MUST FIX", "MUST-FIX"]):
        return "CRITICAL"
    if any(k in upper for k in ["HIGH", "SHOULD FIX", "SHOULD-FIX"]):
        return "HIGH"
    if any(k in upper for k in ["MEDIUM", "NICE TO FIX", "MODERATE"]):
        return "MEDIUM"
    if any(k in upper for k in ["LOW", "BACKLOG", "MINOR"]):
        return "LOW"
    return None


def _build_rec_table(doc, lines):
    """Render a colour-coded recommendation table from LLM output lines."""
    import re
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
    for i, h in enumerate(["Priority", "Finding", "Recommendation"]):
        c = tbl.rows[0].cells[i]
        _ct(c, h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(c, MID_BLUE)

    for line in lines:
        if not line.strip():
            continue
        priority = _detect_priority(line) or "MEDIUM"
        cfg = PRIORITY_CONFIG[priority]
        badge, finding, fix = cfg[2], line.strip(), ""

        # Try to split "finding: fix" or "finding — fix"
        for sep in [" — ", " - ", ": "]:
            if sep in line:
                parts = line.split(sep, 1)
                finding = parts[0].strip()
                fix     = parts[1].strip() if len(parts) > 1 else ""
                break

        tr = tbl.add_row()
        # Priority badge cell
        _ct(tr.cells[0], badge, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(tr.cells[0], cfg[0])
        tr.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(cfg[1])
        # Finding
        _ct(tr.cells[1], finding)
        _shd(tr.cells[1], GREY)
        # Fix
        _ct(tr.cells[2], fix)
        _shd(tr.cells[2], WHITE)


def _build_recommendations(battery, network, snaps, cold, warm):
    """Build structured recommendation list from all test data."""
    recs = []

    # GPS
    gps_act  = battery.get("gps_activations", 0)
    dur_min  = _parse_duration_to_min(battery.get("session_duration",""))
    gps_rate = round(gps_act / dur_min, 1) if dur_min > 0 else 0
    if gps_rate > 6:
        recs.append(("CRITICAL",
            "GPS Over-polling detected",
            f"GPS activates {gps_rate}x/min (standard: <6/min). "
            "Use FusedLocationProviderClient with minInterval=10000ms. "
            "Switch to PRIORITY_BALANCED_POWER when speed <5km/h."))

    # Camera double-open
    cam = battery.get("camera_sessions", 0)
    if cam > 20:
        recs.append(("HIGH",
            f"Camera double-open pattern ({cam} sessions detected)",
            "App opens camera hardware twice per photo capture. "
            "Review CameraManager lifecycle — call unbindAll() after each capture. "
            "Expected sessions = trips x photos x 1 (not x2)."))

    # GL memory growth
    if len(snaps) >= 2:
        gl_first = snaps[0].get("parsed",{}).get("gl_mtrack_mb", 0) or 0
        gl_last  = snaps[-1].get("parsed",{}).get("gl_mtrack_mb", 0) or 0
        if gl_last > gl_first * 1.5 and gl_last > 30:
            recs.append(("HIGH",
                f"GL texture memory not releasing ({gl_first} MB -> {gl_last} MB)",
                "Camera GL textures accumulate across trips. "
                "Call camera.unbindAll() explicitly after each photo. "
                "Add lifecycle-aware cleanup in onPause/onStop."))

    # Memory growth
    if len(snaps) >= 2:
        pss_first = snaps[0].get("parsed",{}).get("total_pss_mb", 0) or 0
        pss_last  = snaps[-1].get("parsed",{}).get("total_pss_mb", 0) or 0
        growth_pct = round(((pss_last - pss_first) / pss_first) * 100) if pss_first else 0
        if growth_pct > 50:
            recs.append(("MEDIUM",
                f"Memory growth across session ({pss_first} MB -> {pss_last} MB, +{growth_pct}%)",
                "Profile with Android Studio Memory Profiler. "
                "Implement explicit cleanup between trip flows. "
                "Acceptable for 1-2 trips/day — monitor if trip count increases."))

    # Play Core double-call
    for rc in network.get("redundant_calls", []):
        recs.append(("LOW",
            rc.get("type", "Redundant network call"),
            rc.get("fix", "Consolidate to single call") +
            f" Evidence: {rc.get('evidence','')}"))

    # Cold start
    cold_avg = cold.get("avg", 0)
    if cold_avg > 2000:
        recs.append(("HIGH",
            f"Cold Start {cold_avg}ms exceeds Google 2000ms threshold",
            "Profile startup with Android Studio App Startup. "
            "Defer non-critical SDK initialisation to background thread."))

    # Swap
    if snaps:
        swap = snaps[0].get("parsed",{}).get("swap_pss_mb", 0) or 0
        if swap > 20:
            recs.append(("HIGH",
                f"Swap memory in use ({swap} MB)",
                "Device is under memory pressure. "
                "Reduce memory footprint or test on higher-RAM device."))

    # Battery test note
    if battery.get("was_charging"):
        recs.append(("MEDIUM",
            "Battery drain not measured — device was charging during test",
            "Re-run with USB unplugged using wireless adb. "
            "See Section 9 for exact steps."))

    # WorkManager
    wm = battery.get("workmanager_jobs", 0)
    if wm > 100:
        recs.append(("LOW",
            f"WorkManager cache not pruned ({wm} job executions recorded)",
            "Call WorkManager.getInstance(context).pruneWork() after trip completion. "
            "Add periodic daily pruning worker."))

    # New Relic detected
    if "New Relic" in network.get("sdks_detected", []):
        recs.append(("LOW",
            "New Relic SDK detected in staging build",
            "Confirm New Relic is intentional in staging. "
            "If not needed, remove to reduce background network overhead."))

    return recs


def _merge_recommendations(llm_recs_raw: list, rule_recs: list) -> list:
    """
    Merge LLM-structured recommendations with rule-based ones.

    llm_recs_raw : list of dicts from analyse_with_llm_structured()
                   each dict has: priority, area, issue, fix,
                                  benchmark_context, estimated_impact
    rule_recs    : list of (priority, issue, fix) tuples from _build_recommendations()

    Returns a unified list of (priority, issue, fix) tuples sorted CRITICAL→LOW.
    LLM recs take precedence; rule-based recs are only added if they don't overlap
    with an existing LLM rec for the same area.
    """
    order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
    merged = []

    # ── 1. Convert and add LLM recs ──────────────────────────────────────
    llm_areas_covered = set()
    for rec in llm_recs_raw:
        priority = str(rec.get("priority", "MEDIUM")).upper()
        if priority not in order:
            priority = "MEDIUM"
        area     = rec.get("area", "Other")
        issue    = rec.get("issue", "")
        fix      = rec.get("fix", "")
        bm       = rec.get("benchmark_context", "")
        impact   = rec.get("estimated_impact", "")

        # Enrich issue text with benchmark and impact if provided
        issue_full = issue
        if bm:
            issue_full += f"  [{bm}]"

        fix_full = fix
        if impact:
            fix_full += f"  Estimated saving: {impact}"

        merged.append((priority, f"[{area}] {issue_full}", fix_full))
        llm_areas_covered.add(area.lower())

    # ── 2. Add rule-based recs that aren't already covered ───────────────
    # Map rule-based areas to llm area names for overlap detection
    area_keywords = {
        "gps":    ["gps", "location", "poll"],
        "memory": ["memory", "gl", "heap", "pss", "swap"],
        "cpu":    ["cpu", "start"],
        "battery":["battery", "charging"],
        "network":["network", "call", "workmanager", "cache", "sdk", "new relic"],
    }

    for (r_priority, r_issue, r_fix) in rule_recs:
        issue_lower = r_issue.lower()
        covered = False
        for area, keywords in area_keywords.items():
            if area in llm_areas_covered and any(k in issue_lower for k in keywords):
                covered = True
                break
        if not covered:
            merged.append((r_priority, r_issue, r_fix))

    # ── 3. Sort: CRITICAL → HIGH → MEDIUM → LOW ─────────────────────────
    return sorted(merged, key=lambda x: order.get(x[0], 99))


def _render_rec_table(doc, recs):
    """Render the full colour-coded recommendations table."""
    if not recs:
        _body(doc, "No recommendations generated — all metrics within acceptable range.")
        return

    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"
    for i, h in enumerate(["#", "Priority", "Issue", "Recommended Fix"]):
        c = tbl.rows[0].cells[i]
        _ct(c, h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(c, MID_BLUE)

    # Sort: CRITICAL first, then HIGH, MEDIUM, LOW
    order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
    recs_sorted = sorted(recs, key=lambda x: order.get(x[0], 99))

    for i, (priority, issue, fix) in enumerate(recs_sorted):
        cfg = PRIORITY_CONFIG.get(priority, PRIORITY_CONFIG["MEDIUM"])
        bg  = GREY if i % 2 else WHITE
        tr  = tbl.add_row()

        # # number
        _ct(tr.cells[0], str(i+1), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shd(tr.cells[0], bg)

        # Priority badge — coloured background, coloured text
        _ct(tr.cells[1], cfg[2], bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
            color=cfg[1])
        _shd(tr.cells[1], cfg[0])

        # Issue
        _ct(tr.cells[2], issue)
        _shd(tr.cells[2], bg)

        # Fix
        _ct(tr.cells[3], fix)
        _shd(tr.cells[3], bg)
